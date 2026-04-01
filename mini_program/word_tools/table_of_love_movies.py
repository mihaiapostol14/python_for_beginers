from docx import Document
from docx.shared import Inches

def create_movies_and_cartoons_table(filename="favorite_movies_and_cartoons.docx"):
    """Creates a table in Word with favorite movies and cartoons."""
    document = Document()

    document.add_heading('My Favorite Movies and Cartoons', level=0)

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Type'
    hdr_cells[1].text = 'Title'
    hdr_cells[2].text = 'Comment (optional)'

    while True:
        item_type = input("Enter type (movie/cartoon) or 'done' to finish: ").lower()
        if item_type == 'done':
            break
        elif item_type not in ['movie', 'cartoon']:
            print("Please enter 'movie' or 'cartoon'.")
            continue

        title = input(f"Enter the title of the {item_type}: ")
        comment = input("Enter a comment (optional): ")

        row_cells = table.add_row().cells
        row_cells[0].text = item_type.capitalize()
        row_cells[1].text = title
        row_cells[2].text = comment

    document.save(filename)
    print(f"Table saved to file '{filename}'")

if __name__ == "__main__":
    create_movies_and_cartoons_table()



