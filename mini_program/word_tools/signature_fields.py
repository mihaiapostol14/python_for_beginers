from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class SignatureDocumentGenerator:
    def __init__(self, output_file='Document_with_Signature.docx'):
        self.doc = Document()
        self.output_file = output_file

    def add_signature(self, align_right=False):
        paragraph = self.doc.add_paragraph()
        paragraph.add_run('Signature:').bold = True
        paragraph.add_run(' ____________________________').bold = True

        if align_right:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    def generate(self, count=20):
        for _ in range(count):
            self.add_signature(align_right=True)
            self.add_signature(align_right=False)

    def save(self):
        self.doc.save(self.output_file)


# Usage
if __name__ == "__main__":
    generator = SignatureDocumentGenerator()
    generator.generate(20)
    generator.save()